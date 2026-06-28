"""Generate the Cratis x CatalystIQ STEP-BY-STEP IMPLEMENTATION PLAN Word document."""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x2D, 0x50)
BLUE = RGBColor(0x2E, 0x5A, 0xAC)
GREEN = RGBColor(0x1E, 0x7A, 0x4B)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def _shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tc.append(shd)


def heading(text, size=15, color=NAVY, space_before=14, space_after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, italic=False, color=None, space_after=6, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, header_fill="1F2D50", widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        _shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def phase(num, title, goal, fill="2E5AAC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"PHASE {num}  ·  {title}")
    r.bold = True
    r.font.size = Pt(13.5)
    r.font.color.rgb = NAVY
    g = doc.add_paragraph()
    g.paragraph_format.space_after = Pt(4)
    gr = g.add_run("Goal:  ")
    gr.bold = True
    gr.font.size = Pt(10.5)
    gr.font.color.rgb = BLUE
    gr2 = g.add_run(goal)
    gr2.font.size = Pt(10.5)
    gr2.font.color.rgb = GREY


def steps(rows):
    table(["Step", "Action", "Owner", "Deliverable"], rows,
          header_fill="2E5AAC")


def exit_criteria(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Exit criteria:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREEN
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True


# ── Title ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cratis × CatalystIQ")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Step-by-Step Implementation Plan")
r.font.size = Pt(14)
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Embedding the WhatsApp ordering + own-fleet delivery engine into the Cratis POS")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Prepared by CatalystIQ  ·  catalystiq.ae  ·  info@catalystiq.ae  ·  +971 58 599 7894\n"
                  "Status: Draft for review  ·  Date: 27 June 2026")
mr.font.size = Pt(9)
mr.font.color.rgb = GREY
doc.add_paragraph()

# ── Owner legend + overview ──────────────────────────────────────────────────
heading("Overview & How To Read This Plan")
body("This plan takes the integration from first kickoff to a live, paying pilot and then to "
     "scale. It is organised into 8 phases. Each phase has a goal, a numbered step table with "
     "an owner and a deliverable, and clear exit criteria — do not start the next phase until "
     "the current one's exit criteria are met.")
body("Owner legend:  CIQ = CatalystIQ (us)   ·   CRT = Cratis   ·   Joint = both teams together.",
     color=GREY)
table(
    ["Phase", "Name", "Indicative duration"],
    [
        ["0", "Alignment & Discovery", "Week 1"],
        ["1", "Foundations — Auth, Provisioning & Connectivity", "Week 1–2"],
        ["2", "Menu Synchronisation", "Week 2–3"],
        ["3", "Order Push (Us → Cratis)", "Week 3–4"],
        ["4", "Status Sync (Cratis → Us)", "Week 4"],
        ["5", "End-to-End on WhatsApp", "Week 5"],
        ["6", "UAT & Single-Restaurant Pilot", "Week 6–7"],
        ["7", "Go-Live, Billing & Scale", "Week 8+"],
    ],
    header_fill="1F2D50",
)
body("Durations are indicative and run partly in parallel; the critical path is the menu and "
     "order decisions in Phase 0.", italic=True, color=GREY)

# ── PHASE 0 ──────────────────────────────────────────────────────────────────
phase(0, "Alignment & Discovery", "Agree the integration shape and exchange everything both sides need to build.")
steps([
    ["0.1", "Kickoff call; name 2 technical contacts each; open a shared channel", "Joint", "Contacts + channel"],
    ["0.2", "Walk through the Integration Requirements doc; settle the 9 open decisions", "Joint", "Signed-off decisions"],
    ["0.3", "Decide MENU MASTER: POS-syncs-to-us vs we-digitize-and-push vs hybrid", "Joint", "Menu direction locked"],
    ["0.4", "Decide order intake (we push vs Cratis polls) and status (webhook vs poll)", "Joint", "Flow direction locked"],
    ["0.5", "Cratis shares API docs, sandbox base URLs, sample menu + order payloads, status list", "CRT", "Cratis API spec"],
    ["0.6", "Agree auth both ways + webhook signing scheme; exchange sandbox credentials", "Joint", "Credentials exchanged"],
    ["0.7", "Provision a Cratis sandbox test store with realistic menu", "CRT", "Test store ready"],
])
exit_criteria("All 9 decisions signed off; both teams hold the other's sandbox URLs and credentials; "
              "menu master and order/status directions are locked.")

# ── PHASE 1 ──────────────────────────────────────────────────────────────────
phase(1, "Foundations — Auth, Provisioning & Connectivity",
      "Stand up identity, store mapping and the subscription signal; prove both sides can talk.")
steps([
    ["1.1", "Issue Cratis an API key per store (X-API-Key); confirm scoping + revoke", "CIQ", "Live API key"],
    ["1.2", "Implement our client to authenticate to Cratis (key or OAuth)", "CIQ", "Outbound auth working"],
    ["1.3", "Build store ↔ restaurant mapping; Cratis sends store IDs, we return restaurant_id", "Joint", "Mapping table"],
    ["1.4", "Implement subscription / active-store signal (activate / suspend / cancel)", "Joint", "Entitlement sync"],
    ["1.5", "Verify webhook signing (HMAC) both directions", "Joint", "Signed webhooks verified"],
    ["1.6", "Connectivity smoke test: authenticated ping each way in sandbox", "Joint", "Green smoke test"],
])
exit_criteria("A sandbox store is mapped, marked active, and both sides can make an authenticated, "
              "signed call to the other.")

# ── PHASE 2 ──────────────────────────────────────────────────────────────────
phase(2, "Menu Synchronisation", "Get the Cratis menu into our system so the WhatsApp bot can sell it.")
steps([
    ["2.1", "Build menu ingest (Cratis push) or pull job, per Phase 0 decision", "CIQ", "Menu sync endpoint/job"],
    ["2.2", "Map fields: item ID, dish number, name, price, category, description, availability", "CIQ", "Field mapping"],
    ["2.3", "Map variants / modifiers / combos from a real sample payload", "Joint", "Variant mapping"],
    ["2.4", "Enforce our rules: number + price required; description carries no price", "CIQ", "Validation pass"],
    ["2.5", "Wire real-time availability (out-of-stock reflects in the bot)", "CIQ", "Live availability"],
    ["2.6", "Incremental sync via updated_at so only changes re-sync", "CIQ", "Delta sync"],
    ["2.7", "Activate the test-store menu and verify it renders correctly in WhatsApp", "Joint", "Menu live in bot"],
])
exit_criteria("The sandbox store's menu (including a variant and a modifier) appears correctly in "
              "the WhatsApp bot, prices match, and an out-of-stock toggle in Cratis hides the item.")

# ── PHASE 3 ──────────────────────────────────────────────────────────────────
phase(3, "Order Push (Us → Cratis)", "Make WhatsApp orders land in the Cratis POS reliably.")
steps([
    ["3.1", "Build the order-push client to Cratis create-order endpoint", "CIQ", "Order push client"],
    ["3.2", "Map order payload: items, qty, variants, notes, totals, address, COD, ETA", "CIQ", "Order mapping"],
    ["3.3", "Send idempotency key; handle Cratis accept/reject + reason", "Joint", "Idempotent intake"],
    ["3.4", "Store the returned Cratis order ID to correlate later status updates", "CIQ", "Order correlation"],
    ["3.5", "Handle modifications (before 'ready') and cancellation per Cratis semantics", "Joint", "Modify/cancel flow"],
    ["3.6", "Retry-with-backoff + dead-letter on push failure; alert on dead-letter", "CIQ", "Reliable delivery"],
    ["3.7", "Place a WhatsApp test order and confirm it appears in the Cratis POS", "Joint", "Order visible in POS"],
])
exit_criteria("A WhatsApp order appears in the Cratis POS with correct items, totals and address; "
              "Cratis returns its order ID; a retried push does not duplicate.")

# ── PHASE 4 ──────────────────────────────────────────────────────────────────
phase(4, "Status Sync (Cratis → Us)", "Drive dispatch from kitchen status and report delivery back.")
steps([
    ["4.1", "Build the status-webhook receiver (or poller) for Cratis events", "CIQ", "Status receiver"],
    ["4.2", "Map exact Cratis status values onto our order FSM", "Joint", "Status mapping"],
    ["4.3", "On 'READY' → trigger automatic rider dispatch + batching", "CIQ", "Ready → dispatch"],
    ["4.4", "On accepted/preparing → customer confirmation + prep countdown on WhatsApp", "CIQ", "Customer updates"],
    ["4.5", "Push delivery status back to Cratis (assigned, picked up, delivered, late)", "CIQ", "Delivery status out"],
    ["4.6", "Late-delivery auto-coupon (except disclosed weather delays)", "CIQ", "Coupon rule live"],
])
exit_criteria("Marking an order 'ready' in Cratis automatically dispatches a rider; the customer "
              "sees live updates; Cratis sees the delivery lifecycle through to 'delivered'.")

# ── PHASE 5 ──────────────────────────────────────────────────────────────────
phase(5, "End-to-End on WhatsApp", "Prove the whole journey on a real WhatsApp number, with edge cases.")
steps([
    ["5.1", "Onboard the test restaurant on its own WhatsApp number", "Joint", "WhatsApp connected"],
    ["5.2", "Digitize / sync the menu and switch ordering on", "CIQ", "Store live (sandbox)"],
    ["5.3", "Run the happy path: order → POS → ready → dispatch → delivered → status back", "Joint", "E2E happy path"],
    ["5.4", "Edge cases: out-of-stock, modification, cancellation, late delivery, COD totals", "Joint", "Edge cases pass"],
    ["5.5", "Batching test: two nearby orders batched under the 40-minute promise", "Joint", "Batching verified"],
])
exit_criteria("The complete order-to-door journey works end-to-end on WhatsApp for the test store, "
              "including the main edge cases and order batching.")

# ── PHASE 6 ──────────────────────────────────────────────────────────────────
phase(6, "UAT & Single-Restaurant Pilot", "Validate with Cratis sign-off, then run one real restaurant.")
steps([
    ["6.1", "Joint UAT against an agreed test script; log and fix defects", "Joint", "UAT sign-off"],
    ["6.2", "Select one real pilot restaurant; provision + go live", "Joint", "Pilot live"],
    ["6.3", "Monitor orders, dispatch, SLA and errors daily for the pilot", "CIQ", "Daily monitoring"],
    ["6.4", "Capture real numbers (orders, on-time %, margin saved) to share with Cratis", "CIQ", "Pilot metrics"],
    ["6.5", "Weekly review; iterate on issues and menu/marketing tuning", "Joint", "Review cadence"],
])
exit_criteria("Cratis signs off UAT; one real restaurant runs live on the integration for an agreed "
              "period with stable metrics and no critical defects.")

# ── PHASE 7 ──────────────────────────────────────────────────────────────────
phase(7, "Go-Live, Billing & Scale", "Switch to production, wire subscription billing, and onboard at scale.")
steps([
    ["7.1", "Promote to production credentials + URLs; production smoke test", "Joint", "Production live"],
    ["7.2", "Wire subscription billing — Cratis bills the restaurant; revenue share agreed", "Joint", "Billing live"],
    ["7.3", "Activate the monthly marketing layer (promo images/videos) per restaurant", "CIQ", "Marketing on"],
    ["7.4", "Publish a support runbook + monitoring/alerting + escalation SLAs", "CIQ", "Support runbook"],
    ["7.5", "Build an onboarding playbook so Cratis can add restaurants repeatably", "Joint", "Onboarding playbook"],
    ["7.6", "Roll out to the next batch of restaurants", "Joint", "Scaled rollout"],
])
exit_criteria("The integration is in production, subscriptions are billing through Cratis, support "
              "and monitoring are in place, and new restaurants can be onboarded from a playbook.")

# ── Responsibilities summary ─────────────────────────────────────────────────
heading("Who Does What — At a Glance")
table(
    ["CatalystIQ (CIQ) builds & runs", "Cratis (CRT) provides", "Joint"],
    [
        ["WhatsApp bot, AI ordering, menu sync logic",
         "POS API + docs + credentials",
         "All Phase-0 decisions"],
        ["Order push, status receiver, dispatch, delivery",
         "Store IDs + subscription signal",
         "Field mapping & sample payloads"],
        ["Riders, batching, SLA, coupons, marketing layer",
         "Order-create endpoint + status webhook",
         "UAT, pilot, go-live"],
        ["Monitoring, support runbook, retries/dead-letter",
         "Restaurant billing + revenue share",
         "Onboarding playbook & scale"],
    ],
    header_fill="1F2D50",
)

# ── Risks / dependencies ─────────────────────────────────────────────────────
heading("Key Dependencies & Risks To Manage", color=BLUE)
bullet("the menu-master decision (Phase 0.3) gates all of Phase 2 — settle it first.",
       bold_lead="Critical path: ")
bullet("COD-only and AED-only today; online payment or multi-currency is net-new work if Cratis needs it.",
       bold_lead="Scope risk: ")
bullet("no per-line tax/VAT yet; if Cratis prices are tax-exclusive we must add it before go-live.",
       bold_lead="Scope risk: ")
bullet("an exact, stable list of Cratis order statuses; mapping breaks if these change later.",
       bold_lead="Dependency: ")
bullet("WhatsApp number provisioning per restaurant can take time — start early in Phase 5.",
       bold_lead="Lead time: ")

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run("CatalystIQ  ·  catalystiq.ae  ·  info@catalystiq.ae  ·  +971 58 599 7894")
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

out = r"c:\Users\user\Desktop\Cratis x CatalystIQ - Implementation Plan.docx"
doc.save(out)
print("saved:", out)
