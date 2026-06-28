"""Final meeting overview doc — simple points, NO hyphens, NO em dashes."""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x2D, 0x50)
BLUE = RGBColor(0x2E, 0x5A, 0xAC)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def _shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tc.append(shd)


def heading(text, size=15, color=NAVY, before=16, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, italic=False, color=None, after=6, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    if lead:
        r = p.add_run(lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, fill="1F2D50"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        _shade(hdr[i], fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cratis  x  CatalystIQ")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Integration Overview")
r.font.size = Pt(15)
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Restaurants get their own WhatsApp ordering and delivery, inside the Cratis POS")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY
doc.add_paragraph()

# 1. What this is
heading("What This Is")
body("CatalystIQ runs the full journey of a food order on WhatsApp. The customer orders by "
     "chat, the kitchen receives it, our own riders deliver it, and the customer watches it "
     "live on a map. We want to place this engine inside the Cratis POS, so Cratis restaurants "
     "get ordering and delivery without ever leaving Cratis.")

# 2. How it works
heading("How It Works")
body("It is a two way link between Cratis and CatalystIQ.")
table(
    ["Direction", "What flows"],
    [
        ["Cratis to Us", "The menu: items, prices, what is available"],
        ["Cratis to Us", "Kitchen status, mainly when an order is ready"],
        ["Us to Cratis", "New WhatsApp orders, straight into the POS"],
        ["Us to Cratis", "Delivery status, until the order reaches the customer"],
    ],
    fill="2E5AAC",
)
body("The API key we give Cratis lets Cratis talk to us. We also need a way to talk back to "
     "Cratis. Both sides connect.", italic=True, color=GREY)

# 3. What we need from Cratis
heading("What We Need From Cratis")
bullet("test and live web addresses, plus simple documentation.", lead="API access:  ")
bullet("a way for us to connect to Cratis, using a key or a login.", lead="Connection:  ")
bullet("one store ID for each restaurant, so our records match.", lead="Store IDs:  ")
bullet("a signal that tells us which restaurants are active and paying.", lead="Active stores:  ")
bullet("the menu, with one sample of how Cratis sends it.", lead="Menu:  ")
bullet("a way to send orders into Cratis, and the list of order statuses Cratis uses.", lead="Orders:  ")
bullet("a test store and two technical contacts.", lead="For testing:  ")

# 4. How we build it
heading("How We Build It")
table(
    ["Step", "What happens", "When"],
    [
        ["1. Align", "Agree the plan and exchange access", "Week 1"],
        ["2. Connect", "Set up the secure link and store matching", "Weeks 1 to 2"],
        ["3. Menu", "Bring the Cratis menu into the WhatsApp bot", "Weeks 2 to 3"],
        ["4. Orders", "Send WhatsApp orders into the Cratis POS", "Weeks 3 to 4"],
        ["5. Status", "Kitchen ready starts delivery, status flows back", "Week 4"],
        ["6. Full test", "Run the whole journey on a real WhatsApp number", "Week 5"],
        ["7. Pilot", "Go live with one real restaurant and review results", "Weeks 6 to 7"],
        ["8. Scale", "Production, billing, and onboard more restaurants", "Week 8 onward"],
    ],
    fill="1F2D50",
)

# 5. Key decisions
heading("Key Decisions To Agree First")
bullet("does the Cratis POS own the menu, or do we build it.", lead="Menu owner:  ")
bullet("do we send orders to Cratis, or does Cratis collect them from us.", lead="Order flow:  ")
bullet("how we connect and keep the link secure.", lead="Connection:  ")
bullet("how Cratis handles VAT and currency.", lead="Tax and currency:  ")

# 6. Good to know
heading("Good To Know", color=BLUE)
bullet("today we support cash on delivery and the AED currency. Online payment or other "
       "currencies would be extra work.", lead="Payment:  ")
bullet("we do not handle VAT per item yet. If Cratis prices exclude VAT, we would add it.",
       lead="Tax:  ")

# 7. Next steps
heading("Next Steps")
bullet("Cratis shares the API documentation and a test store.")
bullet("Both teams agree the key decisions above.")
bullet("We begin with the connect step in a safe test setup.")

# Footer
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run("CatalystIQ   ·   catalystiq.ae   ·   info@catalystiq.ae   ·   +971 58 599 7894")
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

out = r"c:\Users\user\Desktop\Cratis x CatalystIQ Integration Overview.docx"
doc.save(out)
print("saved:", out)
