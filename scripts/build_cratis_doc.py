"""Generate the Cratis x CatalystIQ integration-requirements Word document."""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x2D, 0x50)
BLUE = RGBColor(0x2E, 0x5A, 0xAC)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def _shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tc.append(shd)


def heading(text, size=15, color=NAVY, space_before=14, space_after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, italic=False, color=None, space_after=6, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None, header_fill="1F2D50"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        _shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ── Title block ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cratis × CatalystIQ")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Integration Requirements — What We Need From Cratis")
r.font.size = Pt(14)
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WhatsApp Ordering + Own-Fleet Delivery Engine, embedded in the Cratis POS")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Prepared by CatalystIQ  ·  catalystiq.ae  ·  info@catalystiq.ae  ·  +971 58 599 7894\n"
                  "Status: Draft for review  ·  Date: 27 June 2026")
mr.font.size = Pt(9)
mr.font.color.rgb = GREY

doc.add_paragraph()

# ── 1. Purpose ───────────────────────────────────────────────────────────────
heading("1.  Purpose")
body("This document lists everything we need from Cratis to embed CatalystIQ's WhatsApp "
     "ordering and own-fleet delivery engine inside the Cratis POS. Cratis does not use our "
     "manager dashboard — the integration is API-only. Cratis drives everything from within "
     "its own POS using an API key we issue.")
body("Key point: this is not simply \"we give Cratis an API key.\" Because Cratis is the POS "
     "(owns the menu) and the destination for orders (kitchen and billing), the integration is "
     "two-way. The key we issue lets Cratis call us; separately, we need credentials and "
     "endpoints to call Cratis. Both directions are required.", italic=True, color=GREY)

# ── 2. Commercial context ────────────────────────────────────────────────────
heading("2.  Commercial Context (from the Cratis × CatalystIQ deck)")
bullet(" Cratis bills the restaurant a monthly subscription per restaurant; CatalystIQ shares "
       "the recurring revenue.", bold_lead="Billing & revenue. ")
bullet(" The service lives inside the Cratis POS, under the Cratis brand. The restaurant stays "
       "Cratis's customer.", bold_lead="Branding & ownership. ")
bullet(" Each restaurant uses its own WhatsApp number and its own delivery fleet. Cash on "
       "delivery, no per-order commission.", bold_lead="Restaurant setup. ")
bullet(" A monthly marketing layer (promo images and short videos) gives restaurants a reason "
       "to stay subscribed.", bold_lead="Marketing layer. ")
body("Two engineering consequences follow and are captured below: we need a subscription / "
     "active-store signal (Section 5) because Cratis owns billing, and we need a clear answer "
     "on who owns the menu (Section 6).", italic=True, color=GREY)

# ── 3. How it works ──────────────────────────────────────────────────────────
heading("3.  How the Integration Works — Two-Way Data Flow")
body("We are the ordering and delivery brain on WhatsApp. Cratis is the source of truth for "
     "the menu and the destination for orders (kitchen, billing, accounting).")
table(
    ["Direction", "What flows", "Why it matters"],
    [
        ["Cratis → Us", "Menu: items, dish numbers, prices, availability, variants",
         "The WhatsApp bot can only sell what Cratis sends"],
        ["Cratis → Us", "Kitchen status (accepted / preparing / READY)",
         "\"Ready\" is what triggers our automatic rider dispatch"],
        ["Us → Cratis", "Orders placed on WhatsApp",
         "So they appear in the POS, kitchen and billing"],
        ["Us → Cratis", "Delivery status (rider assigned, picked up, delivered, late)",
         "So the POS shows the full order lifecycle"],
    ],
)
body("What we already provide today: API-key issuance and revocation, and a read-only partner "
     "data API. Everything in Section 5 is what we need from Cratis to complete the loop.",
     italic=True, color=GREY)

# ── 4. Two-way auth ──────────────────────────────────────────────────────────
heading("4.  Authentication — Both Directions")
body("There are two authentication directions. Please confirm both.")
bullet(" We issue Cratis an API key per store, sent as the header X-API-Key. Shown once at "
       "creation, scoped to one store, revocable.", bold_lead="4a. Cratis → Us. ")
bullet(" We need credentials to push orders into Cratis and to verify webhooks Cratis sends "
       "us. Please confirm: API key or OAuth 2.0 client credentials, plus your webhook signing "
       "scheme (e.g. HMAC shared secret in a header).",
       bold_lead="4b. Us → Cratis. ")
body("Security questions: IP allow-listing on either side? TLS / certificate requirements? "
     "PII constraints on customer name, phone and address?", color=GREY)

# ── 5. The full ask list ─────────────────────────────────────────────────────
heading("5.  What We Need From Cratis — The Complete List")
table(
    ["#", "What we need", "Notes"],
    [
        ["1", "Sandbox + production API base URLs and documentation",
         "OpenAPI/Swagger or equivalent for the endpoints we will call"],
        ["2", "Credentials for us to call Cratis (key or OAuth) + webhook signing",
         "The key we give you is not enough — we must authenticate to you too"],
        ["3", "Store / branch IDs to map to our restaurant IDs",
         "One stable ID per location; we return our restaurant_id"],
        ["4", "Subscription / active-store status per store  (NEW — from the deck)",
         "Cratis owns billing, so Cratis tells us which stores are live/suspended"],
        ["5", "Menu decision + sample payload  (who is master?)",
         "POS menu syncs to us, or we digitize and push into Cratis — see Section 6"],
        ["6", "Cratis order-create endpoint + sample payload",
         "So WhatsApp orders land in the POS"],
        ["7", "Exact Cratis order status values + a status webhook",
         "We map them to our delivery FSM; we never invent statuses"],
        ["8", "Tax/VAT, currency, and delivery-fee ownership answers",
         "See Section 9"],
        ["9", "A test store with realistic data + 2 technical contacts",
         "For end-to-end sandbox integration and UAT"],
    ],
    header_fill="2E5AAC",
)

# ── 6. Menu ──────────────────────────────────────────────────────────────────
heading("6.  Menu — The Big Open Question: Who Is Master?")
body("The deck says \"menu digitized in a day\" — that is our capability (we scan a menu and "
     "build it). But a POS already holds the menu. So which is the source of truth? This "
     "decides the direction of the menu integration and must be settled first.")
table(
    ["Option", "Meaning"],
    [
        ["A — POS is master", "Cratis POS menu syncs to us (Cratis pushes, or we pull on a schedule)"],
        ["B — We are master", "We digitize the menu and push it into Cratis"],
        ["C — Hybrid", "We digitize once, then the Cratis POS owns it going forward"],
    ],
)
body("Required fields per menu item (our model is dish-number + price driven):")
table(
    ["Field", "Required", "Notes"],
    [
        ["Cratis item ID (stable)", "Yes", "So updates map to the same dish"],
        ["Dish number", "Yes", "Mandatory — menu cannot go live if any item lacks a number"],
        ["Name", "Yes", ""],
        ["Price", "Yes", "Decimal, 2 dp — no item activates without a price"],
        ["Category", "Optional", "e.g. Biryani, Drinks — for grouping"],
        ["Description", "Optional", "Customer-facing. Max 3 lines, must NOT contain price"],
        ["Availability (in/out of stock)", "Yes", "Drives \"sold out\" in the bot in real time"],
        ["Variants / sizes (name + price)", "Optional", "e.g. Small/Large — see note"],
        ["Prep time (minutes)", "Optional", "Improves SLA and kitchen countdown accuracy"],
    ],
)
body("Variants / modifiers / combos: please share a sample item with a variant and a priced "
     "modifier so we can lock the mapping. Our hard rules: every sellable item needs a number "
     "and a price; customer-facing descriptions carry no price; COD only.", color=GREY)

# ── 7. Orders ────────────────────────────────────────────────────────────────
heading("7.  Orders — What We Push Into Cratis")
body("When a customer completes an order on WhatsApp, it must land in Cratis. Decision: we "
     "POST the order to a Cratis create-order endpoint (preferred), or Cratis polls our partner "
     "order API. Fields we will send:")
table(
    ["Field", "Notes"],
    [
        ["Our order number", "Stable reference, e.g. R1-0042"],
        ["Cratis store ID", "Which branch (from Section 5 mapping)"],
        ["Customer name + phone", "From the WhatsApp conversation"],
        ["Line items", "dish number, name, variant, qty, price, notes"],
        ["Item notes / special requests", "Verbatim, e.g. \"no onion\""],
        ["Subtotal, delivery fee, total", "Delivery fee tiered by distance (AED)"],
        ["Payment method", "COD only (today)"],
        ["Delivery address", "building, room/apt, receiver name, lat/lng, extra details"],
        ["Promised ETA / SLA deadline", "Customer is told 40 minutes"],
        ["Order placed time", "UTC"],
    ],
)
body("What Cratis must return on intake: the Cratis order ID (so we can correlate status), and "
     "accept/reject with a reason. Modifications are allowed only before \"ready\" — please tell "
     "us whether to send a full replacement or a delta, and how Cratis represents cancellation "
     "and auto-resale.", color=GREY)

# ── 8. Status ────────────────────────────────────────────────────────────────
heading("8.  Order Status — What Cratis Sends Back")
body("Our delivery engine is driven by kitchen status. We need a webhook from Cratis (or a "
     "pollable status endpoint) on every change.")
table(
    ["Cratis event", "What it triggers on our side"],
    [
        ["Order accepted by store", "Confirm to the customer on WhatsApp"],
        ["Preparing / in kitchen", "Start the prep countdown"],
        ["READY", "Triggers automatic rider dispatch + batching (critical)"],
        ["Rejected / cancelled by store", "Notify the customer and release"],
    ],
)
body("We need the exact list of Cratis status values and their meaning. In return we send back: "
     "rider assigned → picked up → en route → delivered, plus a late flag and automatic coupon "
     "on late delivery (except disclosed weather delays).", color=GREY)

# ── 9. Money/tax/time ────────────────────────────────────────────────────────
heading("9.  Money, Tax, Time & Locale")
table(
    ["Topic", "Our assumption today", "Need Cratis to confirm"],
    [
        ["Currency", "AED only", "Multi-currency? Per-store currency code?"],
        ["Decimals", "2 dp", "Match?"],
        ["Tax / VAT", "Not modelled per line yet", "Prices tax-inclusive or exclusive? VAT line needed?"],
        ["Delivery fee", "We compute (distance tiers)", "Should Cratis own the fee, or accept ours?"],
        ["Timezone", "Asia/Dubai display, UTC storage", "Confirm per store"],
        ["Payment", "COD only", "Online payment via Cratis later?"],
    ],
)

# ── 10. Gaps on our side ─────────────────────────────────────────────────────
heading("10.  Gaps On Our Side To Flag Early", color=BLUE)
bullet(" We are COD-only and AED-only today. If Cratis wants online payment or multi-currency, "
       "that is net-new work on our side.", bold_lead="Payment & currency. ")
bullet(" We do not model tax/VAT per line yet. If Cratis prices are tax-exclusive, we will need "
       "to add it.", bold_lead="Tax/VAT. ")

# ── 11. Decisions ────────────────────────────────────────────────────────────
heading("11.  Open Decisions To Settle Jointly")
for i, d in enumerate([
    "Menu master — POS syncs to us, or we digitize and push into Cratis?",
    "Menu sync mechanism — Cratis pushes to us, or we pull?",
    "Order intake — we push to Cratis, or Cratis polls us?",
    "Status updates — Cratis webhook to us, or we poll?",
    "Auth direction Us → Cratis — API key or OAuth? Webhook signing scheme?",
    "Variants / modifiers / combos representation.",
    "Tax/VAT handling and who owns the delivery fee.",
    "Subscription/active-store signal — webhook or status field?",
    "Multi-currency / multi-store scope for phase 1.",
], 1):
    bullet(d, bold_lead=f"{i}.  ")

# ── 12. To start ─────────────────────────────────────────────────────────────
heading("12.  What We Need To Receive To Start Building")
for item in [
    "Cratis API docs (sandbox) and base URLs.",
    "Credentials for us to call Cratis (Section 4b).",
    "A sample menu payload (including a variant and a modifier).",
    "A sample order payload Cratis expects, or the spec of your create-order endpoint.",
    "The list of Cratis order status values.",
    "A test store with realistic data and two technical contacts.",
]:
    bullet(item)
body("Once we have these, we can map fields, stand up the sandbox integration, and schedule a "
     "joint UAT.", italic=True, color=GREY)

# ── Footer line ──────────────────────────────────────────────────────────────
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run("CatalystIQ  ·  catalystiq.ae  ·  info@catalystiq.ae  ·  +971 58 599 7894")
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

out = r"c:\Users\user\Desktop\Cratis x CatalystIQ - Integration Requirements.docx"
doc.save(out)
print("saved:", out)
